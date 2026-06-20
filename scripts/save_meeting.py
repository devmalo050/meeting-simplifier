#!/usr/bin/env python3
# scripts/save_meeting.py
# 사용법: python save_meeting.py --title TITLE --minutes-file PATH --audio-path PATH --format md --output-dir ~/Documents/meetings
# JSON 결과 stdout 출력
import sys
import os
import json
import re
import shutil
import argparse
from datetime import datetime
from pathlib import Path

import config


def sanitize_dir_name(title):
    cleaned = re.sub(r'[<>:"/\\|?*]', '', title)
    cleaned = re.sub(r'\s+', '-', cleaned)
    return cleaned[:80]


def save_docx(file_path, title, minutes):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx가 설치되지 않았습니다. setup.sh를 실행하세요.")

    doc = Document()
    for line in minutes.split('\n'):
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    doc.save(file_path)


def save_meeting(title, minutes, audio_path, fmt, output_dir, transcript_file=''):
    # 전체 트랜스크립트는 LLM 본문이 아니라 파일에서 직접 붙인다 (길어도 잘리지 않도록)
    if transcript_file and os.path.exists(transcript_file):
        with open(transcript_file, 'r', encoding='utf-8') as tf:
            transcript_text = tf.read().strip()
        if transcript_text:
            minutes = minutes.rstrip() + "\n\n## 전체 트랜스크립트\n\n" + transcript_text + "\n"

    resolved_base = str(Path(output_dir).expanduser())
    now = datetime.now()
    date_str = now.strftime('%Y-%m-%d')
    time_str = now.strftime('%H%M%S')
    safe_title = sanitize_dir_name(title)
    dir_name = sanitize_dir_name(f"{date_str}-{time_str}-{title}")
    meeting_dir = os.path.join(resolved_base, dir_name)

    try:
        os.makedirs(meeting_dir, exist_ok=True)
    except PermissionError:
        fallback_dir = os.path.join(Path.home(), 'Desktop', dir_name)
        os.makedirs(fallback_dir, exist_ok=True)
        meeting_dir = fallback_dir

    audio_moved = False
    audio_error = None
    if audio_path:
        ext = os.path.splitext(audio_path)[1] or '.wav'
        dest_audio = os.path.join(meeting_dir, f"{safe_title}{ext}")
        try:
            shutil.move(audio_path, dest_audio)
            audio_moved = True
        except Exception as e:
            audio_error = str(e)
            sys.stderr.write(f"[save_meeting] 오디오 파일 이동 실패: {e}\n")

    minutes_path = os.path.join(meeting_dir, f"{safe_title}.{fmt}")
    if fmt in ('md', 'txt'):
        with open(minutes_path, 'w', encoding='utf-8') as f:
            f.write(minutes)
    elif fmt == 'docx':
        save_docx(minutes_path, title, minutes)

    result = {"saved_dir": meeting_dir}
    if audio_path:
        result["audio_moved"] = audio_moved
        if not audio_moved:
            result["audio_source"] = audio_path
            result["audio_error"] = audio_error
    return result


def resolve_output_dir(arg_output_dir):
    if arg_output_dir:
        return arg_output_dir
    return config.effective_output_dir()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--title', required=True)
    parser.add_argument('--minutes-file', required=True, help='회의록 내용이 담긴 임시 파일 경로')
    parser.add_argument('--audio-path', default='')
    parser.add_argument('--format', default='md', choices=['md', 'txt', 'docx'])
    parser.add_argument('--output-dir', default=None)
    parser.add_argument('--transcript-file', default='', help='전체 트랜스크립트가 담긴 파일 경로 (본문 끝에 코드로 추가됨)')
    args = parser.parse_args()

    with open(args.minutes_file, 'r', encoding='utf-8') as f:
        minutes = f.read()

    try:
        result = save_meeting(
            title=args.title,
            minutes=minutes,
            audio_path=args.audio_path or '',
            fmt=args.format,
            output_dir=resolve_output_dir(args.output_dir),
            transcript_file=args.transcript_file or '',
        )
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"error": str(e)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == '__main__':
    main()
